from __future__ import annotations

import json
import os
import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import RGBColor, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from playwright.sync_api import sync_playwright
from reportlab.lib.colors import HexColor
from reportlab.pdfgen import canvas


API_BASE = "http://127.0.0.1:8765"
APP_BASE = "http://127.0.0.1:5173"


def data(response):
    assert response.ok, f"{response.status}: {response.text()}"
    return response.json()["data"]


def pdf_bytes() -> bytes:
    stream = BytesIO()
    page = canvas.Canvas(stream)
    page.setFillColor(HexColor("#2F5597"))
    page.setFont("Helvetica-Bold", 24)
    page.drawString(72, 760, "Rich PDF layout")
    page.rect(72, 680, 300, 52, fill=0, stroke=1)
    page.save()
    return stream.getvalue()


def docx_bytes() -> bytes:
    stream = BytesIO()
    document = Document()
    title = document.add_heading("Rich Word layout", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x2F, 0x55, 0x97)
    title.runs[0].font.size = Pt(26)
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Accuracy"
    table.cell(1, 1).text = "95%"
    document.save(stream)
    return stream.getvalue()


def xlsx_bytes() -> bytes:
    stream = BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Styled"
    sheet.merge_cells("A1:B1")
    sheet["A1"] = "Quarterly report"
    sheet["A1"].font = Font(size=16, bold=True, color="FFFFFF")
    sheet["A1"].fill = PatternFill(fill_type="solid", fgColor="2F5597")
    sheet["A1"].alignment = Alignment(horizontal="center", vertical="center")
    sheet["A2"] = "Revenue"
    sheet["B2"] = 125000
    sheet.column_dimensions["A"].width = 24
    sheet.row_dimensions[1].height = 30
    workbook.save(stream)
    return stream.getvalue()


def notebook_bytes() -> bytes:
    return json.dumps(
        {
            "nbformat": 4,
            "nbformat_minor": 5,
            "metadata": {"kernelspec": {"name": "python3", "display_name": "Python 3"}},
            "cells": [
                {"cell_type": "markdown", "metadata": {}, "source": ["# Notebook layout\n", "Rendered **Markdown** cell."]},
                {
                    "cell_type": "code",
                    "execution_count": 3,
                    "metadata": {},
                    "source": ["score = 0.95\n", "score"],
                    "outputs": [{"output_type": "execute_result", "execution_count": 3, "metadata": {}, "data": {"text/plain": ["0.95"]}}],
                },
            ],
        },
        ensure_ascii=False,
    ).encode("utf-8")


with sync_playwright() as playwright:
    browser_root = Path(os.environ["LOCALAPPDATA"]) / "ms-playwright"
    installed_shells = sorted(browser_root.glob("chromium_headless_shell-*/chrome-headless-shell-win64/chrome-headless-shell.exe"))
    browser = playwright.chromium.launch(headless=True, executable_path=str(installed_shells[-1]) if installed_shells else None)
    context = browser.new_context(viewport={"width": 1600, "height": 1000})
    page = context.new_page()
    console_errors: list[str] = []
    page.on("console", lambda message: console_errors.append(message.text) if message.type == "error" else None)
    page.add_init_script(
        f"""
        window.studypilot = {{
          runtime: async () => ({{ apiBase: {json.dumps(API_BASE)}, dataDir: 'test-data' }}),
          window: {{ minimize() {{}}, toggleMaximize() {{}}, close() {{}} }},
          clipboard: {{ readText: async () => '', readImage: async () => null }},
          files: {{ saveExport: async () => null }}
        }};
        """
    )

    courses = data(page.request.get(f"{API_BASE}/api/courses"))
    course_id = courses[0]["id"]
    data(page.request.put(f"{API_BASE}/api/settings/onboarding_complete", data={"value": True}))

    fixtures = {}
    for position, (filename, mime_type, content) in enumerate([
        ("rich-layout.pdf", "application/pdf", pdf_bytes()),
        ("rich-layout.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", docx_bytes()),
        ("rich-layout.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", xlsx_bytes()),
        ("rich-layout.ipynb", "application/x-ipynb+json", notebook_bytes()),
    ]):
        fixtures[filename] = data(page.request.post(
            f"{API_BASE}/api/documents/import",
            multipart={
                "source_created_at": f"2025-01-0{position + 1}T04:05:06.000Z",
                "file": {"name": filename, "mimeType": mime_type, "buffer": content},
            },
        ))

    page.goto(f"{APP_BASE}/courses/{course_id}/library")
    page.wait_for_load_state("networkidle")
    books = page.locator(".document-book")
    assert books.count() == 4
    first_book = books.first
    first_box = first_book.bounding_box()
    assert first_box and first_box["height"] > first_box["width"]
    assert page.locator(".document-book__spine").count() == 4
    assert page.get_by_text("文件创建", exact=True).count() == 4
    assert page.get_by_text("导入时间", exact=True).count() == 4
    page.get_by_label("筛选资料格式").select_option("xlsx")
    assert books.count() == 1
    assert "1 / 4" in page.locator(".library-controls__count").inner_text()
    page.get_by_label("筛选资料格式").select_option("all")
    page.screenshot(path="artifacts/document-bookshelf.png", full_page=True)

    pdf_id = fixtures["rich-layout.pdf"]["id"]
    page.goto(f"{APP_BASE}/courses/{course_id}/library/documents/{pdf_id}")
    page.wait_for_load_state("networkidle")
    pdf_frame = page.get_by_title("PDF 原版阅读器")
    assert pdf_frame.is_visible()
    assert f"/api/documents/{pdf_id}/file" in pdf_frame.get_attribute("src")
    raw_pdf = page.request.get(f"{API_BASE}/api/documents/{pdf_id}/file")
    assert raw_pdf.ok and raw_pdf.headers["content-type"].startswith("application/pdf")

    word_id = fixtures["rich-layout.docx"]["id"]
    page.goto(f"{APP_BASE}/courses/{course_id}/library/documents/{word_id}")
    page.get_by_role("region", name="Word 原版阅读器").wait_for()
    page.locator(".word-reader-shell[data-state='ready']").wait_for(timeout=15_000)
    assert page.locator(".word-layout-reader").get_by_text("Rich Word layout").is_visible()
    assert page.locator(".word-layout-reader table").is_visible()

    xlsx_id = fixtures["rich-layout.xlsx"]["id"]
    page.goto(f"{APP_BASE}/courses/{course_id}/library/documents/{xlsx_id}")
    spreadsheet_cell = page.get_by_role("gridcell", name=re.compile(r"A1 Quarterly report"))
    spreadsheet_cell.wait_for()
    cell_background = spreadsheet_cell.evaluate("element => getComputedStyle(element).backgroundColor")
    assert spreadsheet_cell.evaluate("element => element.style.gridColumn").endswith("span 2")
    assert cell_background == "rgb(47, 85, 151)"
    page.screenshot(path="artifacts/rich-xlsx-reader.png", full_page=True)

    ipynb_id = fixtures["rich-layout.ipynb"]["id"]
    page.goto(f"{APP_BASE}/courses/{course_id}/library/documents/{ipynb_id}")
    page.get_by_role("heading", name="Notebook layout").wait_for()
    assert page.get_by_text("In [3]").is_visible()
    assert page.get_by_text("0.95", exact=True).is_visible()

    notebooks = data(page.request.get(f"{API_BASE}/api/courses/{course_id}/notebooks"))
    notebook_id = notebooks[0]["id"]
    page.goto(f"{APP_BASE}/courses/{course_id}/knowledge/{notebook_id}")
    page.wait_for_load_state("networkidle")
    page.get_by_role("button", name=re.compile("PILOT")).first.click()
    page.get_by_role("complementary", name=re.compile("PILOT")).wait_for()
    page.wait_for_timeout(450)
    toolbar = page.locator(".canvas-toolbar")
    secondary = page.locator(".canvas-toolbar__secondary")
    nav = page.locator(".tool-group--canvas-nav")
    toolbar_box = toolbar.bounding_box()
    secondary_box = secondary.bounding_box()
    assert toolbar.evaluate("element => getComputedStyle(element).display") == "grid"
    assert toolbar_box and secondary_box and secondary_box["y"] > toolbar_box["y"]
    assert secondary_box["y"] + secondary_box["height"] <= toolbar_box["y"] + toolbar_box["height"] + 2
    assert nav.evaluate("element => getComputedStyle(element).overflowX") == "visible"
    assert nav.evaluate("element => element.scrollWidth <= element.clientWidth + 2")
    page.screenshot(path="artifacts/knowledge-toolbar-agent-open.png", full_page=True)

    assert console_errors == [], console_errors
    print(json.dumps({
        "pdf_original": True,
        "word_layout": True,
        "xlsx_styles": True,
        "ipynb_cells": True,
        "book_library": True,
        "knowledge_toolbar_reflow": True,
        "console_errors": len(console_errors),
    }, ensure_ascii=False))
    browser.close()
