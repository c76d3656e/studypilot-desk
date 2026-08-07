from __future__ import annotations

import base64
import math
import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from ..db import Database
from ..errors import AppError
from ..repository import as_dict


@dataclass(frozen=True)
class ExportArtifact:
    content: bytes
    filename: str
    media_type: str


KIND_LABELS = {
    "concept": "概念",
    "sticky_note": "便签",
    "flashcard": "记忆卡",
    "citation": "资料摘录",
    "image": "图片",
}
RELATION_LABELS = {
    "prerequisite": "前置依赖",
    "mindmap": "思维分支",
    "association": "自由关联",
}
CARD_COLORS = {
    "indigo": (91, 101, 225),
    "sun": (207, 142, 49),
    "mint": (36, 158, 134),
    "coral": (205, 89, 101),
    "slate": (91, 103, 122),
    "blue": (69, 111, 219),
    "teal": (35, 143, 137),
    "yellow": (202, 143, 42),
}


class NotebookExportService:
    def __init__(self, database: Database, data_dir: Path) -> None:
        self.database = database
        self.media_root = (data_dir / "media").resolve()

    def export(
        self,
        course_id: int,
        notebook_id: int,
        format_name: str,
        canvas_width: int = 1800,
        canvas_height: int = 1100,
    ) -> ExportArtifact:
        course, notebook, nodes, edges = self._read_graph(course_id, notebook_id)
        ordered = sorted(
            nodes,
            key=lambda node: (
                float(node.get("position_y") or 0),
                float(node.get("position_x") or 0),
                int(node["id"]),
            ),
        )
        safe_title = self._safe_filename(notebook["title"])
        if format_name == "png":
            content = self._render_canvas(nodes, edges, canvas_width, canvas_height, "PNG")
            return ExportArtifact(content, f"{safe_title}-知识画布.png", "image/png")
        if format_name == "pdf":
            content = self._render_canvas(nodes, edges, canvas_width, canvas_height, "PDF")
            return ExportArtifact(content, f"{safe_title}-知识画布.pdf", "application/pdf")
        if format_name == "docx":
            content = self._render_docx(course, notebook, ordered, edges)
            return ExportArtifact(
                content,
                f"{safe_title}-知识笔记.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        if format_name == "md":
            content = self._render_markdown(course, notebook, ordered, edges).encode("utf-8")
            return ExportArtifact(content, f"{safe_title}-知识笔记.md", "text/markdown; charset=utf-8")
        raise AppError("UNSUPPORTED_EXPORT", "不支持的导出格式", 422)

    def _read_graph(self, course_id: int, notebook_id: int):
        with self.database.connect() as connection:
            notebook_row = connection.execute(
                """SELECT notebooks.*, courses.title AS course_title
                FROM knowledge_notebooks AS notebooks
                JOIN courses ON courses.id=notebooks.course_id
                WHERE notebooks.id=? AND notebooks.course_id=?
                  AND notebooks.deleted_at IS NULL AND courses.deleted_at IS NULL""",
                (notebook_id, course_id),
            ).fetchone()
            if not notebook_row:
                raise AppError("NOTEBOOK_NOT_FOUND", "知识笔记不存在", 404)
            node_rows = connection.execute(
                """SELECT nodes.*, assets.filename AS image_filename,
                    assets.media_type AS image_media_type,
                    assets.storage_path AS image_storage_path
                FROM knowledge_nodes AS nodes
                LEFT JOIN media_assets AS assets
                  ON assets.id=nodes.image_asset_id AND assets.course_id=nodes.course_id
                WHERE nodes.course_id=? AND nodes.notebook_id=? ORDER BY nodes.id""",
                (course_id, notebook_id),
            ).fetchall()
            edge_rows = connection.execute(
                """SELECT * FROM knowledge_edges
                WHERE course_id=? AND notebook_id=? ORDER BY id""",
                (course_id, notebook_id),
            ).fetchall()
        notebook = as_dict(notebook_row)
        course = {"id": course_id, "title": notebook.pop("course_title")}
        return course, notebook, [as_dict(row) for row in node_rows], [as_dict(row) for row in edge_rows]

    def _render_canvas(
        self,
        nodes: list[dict],
        edges: list[dict],
        requested_width: int,
        requested_height: int,
        target: str,
    ) -> bytes:
        card_sizes = {int(node["id"]): self._card_size(node) for node in nodes}
        content_width = max(
            [1200, requested_width, *[int(float(node.get("position_x") or 0)) + card_sizes[int(node["id"])][0] + 80 for node in nodes]],
        )
        content_height = max(
            [800, requested_height, *[int(float(node.get("position_y") or 0)) + card_sizes[int(node["id"])][1] + 80 for node in nodes]],
        )
        width, height = min(content_width, 4200), min(content_height, 2800)
        image = Image.new("RGB", (width, height), (246, 247, 251))
        draw = ImageDraw.Draw(image)
        for y in range(20, height, 28):
            for x in range(20, width, 28):
                draw.ellipse((x, y, x + 2, y + 2), fill=(215, 219, 230))

        node_by_id = {int(node["id"]): node for node in nodes}
        for edge in edges:
            source = node_by_id.get(int(edge["source_id"]))
            target_node = node_by_id.get(int(edge["target_id"]))
            if not source or not target_node:
                continue
            sw, sh = card_sizes[int(source["id"])]
            tw, th = card_sizes[int(target_node["id"])]
            start = (int(float(source.get("position_x") or 0)) + sw, int(float(source.get("position_y") or 0)) + sh // 2)
            end = (int(float(target_node.get("position_x") or 0)), int(float(target_node.get("position_y") or 0)) + th // 2)
            color = (125, 133, 158)
            draw.line((start, end), fill=color, width=3)
            self._draw_arrow(draw, start, end, color)

        for node in nodes:
            self._draw_card(image, draw, node, card_sizes[int(node["id"])])

        stream = BytesIO()
        if target == "PNG":
            image.save(stream, format="PNG", optimize=True)
        else:
            image.save(stream, format="PDF", resolution=144.0, title="StudyPilot 知识画布")
        return stream.getvalue()

    def _draw_card(self, canvas: Image.Image, draw: ImageDraw.ImageDraw, node: dict, size: tuple[int, int]) -> None:
        x = max(20, int(float(node.get("position_x") or 0)))
        y = max(20, int(float(node.get("position_y") or 0)))
        width, height = size
        accent = CARD_COLORS.get(str(node.get("color")), CARD_COLORS.get(str(node.get("kind")), (91, 101, 225)))
        draw.rounded_rectangle((x + 7, y + 9, x + width + 7, y + height + 9), radius=18, fill=(218, 221, 229))
        draw.rounded_rectangle((x, y, x + width, y + height), radius=18, fill=(255, 255, 255), outline=(210, 214, 225), width=2)
        draw.rounded_rectangle((x, y, x + 8, y + height), radius=7, fill=accent)
        font_scale = min(2.0, max(0.7, float(node.get("font_scale") or 1.0)))
        label_font = self._font(max(10, round(13 * font_scale)), bold=True)
        title_font = self._font(max(15, round(21 * font_scale)), bold=True)
        body_font = self._font(max(11, round(15 * font_scale)))
        label = KIND_LABELS.get(str(node.get("kind")), "知识卡")
        draw.rounded_rectangle((x + 20, y + 18, x + 84, y + 44), radius=8, fill=tuple(min(255, channel + 150) for channel in accent))
        draw.text((x + 31, y + 23), label, fill=accent, font=label_font)
        module = str(node.get("module") or "未分类")
        draw.text((x + width - min(width - 110, self._text_width(draw, module, label_font)) - 18, y + 24), module, fill=(119, 126, 145), font=label_font)
        cursor_y = y + 57
        for line in self._wrap(draw, str(node.get("title") or "未命名卡片"), title_font, width - 42)[:2]:
            draw.text((x + 20, cursor_y), line, fill=(27, 31, 43), font=title_font)
            cursor_y += 28
        image_path = self._image_path(node)
        if image_path:
            try:
                with Image.open(image_path) as source:
                    source = source.convert("RGB")
                    source.thumbnail((width - 40, max(80, min(260, height - cursor_y + y - 42))))
                    image_x = x + (width - source.width) // 2
                    canvas.paste(source, (image_x, cursor_y + 4))
                    cursor_y += source.height + 16
            except OSError:
                pass
        body = str(node.get("content") or node.get("description") or "")
        line_height = max(17, round(22 * font_scale))
        available_lines = max(1, (y + height - 18 - cursor_y) // line_height)
        for line in self._wrap(draw, body, body_font, width - 42)[:available_lines]:
            draw.text((x + 20, cursor_y), line, fill=(82, 89, 107), font=body_font)
            cursor_y += line_height

    def _render_docx(self, course: dict, notebook: dict, nodes: list[dict], edges: list[dict]) -> bytes:
        document = Document()
        section = document.sections[0]
        section.start_type = WD_SECTION.NEW_PAGE
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.1)
        section.bottom_margin = Cm(2.1)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        normal = document.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = RGBColor(62, 68, 82)
        title = document.add_heading(notebook["title"], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta = document.add_paragraph(f"课程：{course['title']}  ·  {len(nodes)} 张知识卡  ·  {len(edges)} 条关系")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if notebook.get("description"):
            document.add_paragraph(str(notebook["description"]))
        document.add_heading("画布内容", level=1)
        if not nodes:
            document.add_paragraph("这张画布暂时还没有知识卡片。")
        for index, node in enumerate(nodes, start=1):
            document.add_heading(f"{index}. {node['title']}", level=2)
            document.add_paragraph(
                f"{KIND_LABELS.get(str(node.get('kind')), '知识卡')}  ·  {node.get('module') or '未分类'}",
            )
            body = str(node.get("content") or node.get("description") or "").strip()
            if body:
                document.add_paragraph(body)
            image_path = self._image_path(node)
            if image_path:
                try:
                    document.add_picture(str(image_path), width=Inches(5.7))
                    if node.get("image_alt"):
                        caption = document.add_paragraph(str(node["image_alt"]))
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except (OSError, ValueError):
                    document.add_paragraph("[图片暂时无法嵌入]")
            if node.get("source_title"):
                document.add_paragraph(f"来源：{node['source_title']}")
            if node.get("source_quote"):
                document.add_paragraph(f"摘录：{node['source_quote']}")
        self._append_docx_relations(document, nodes, edges)
        stream = BytesIO()
        document.save(stream)
        return stream.getvalue()

    def _append_docx_relations(self, document: Document, nodes: list[dict], edges: list[dict]) -> None:
        document.add_heading("关系清单", level=1)
        names = {int(node["id"]): str(node["title"]) for node in nodes}
        if not edges:
            document.add_paragraph("尚未建立卡片关系。")
            return
        for edge in edges:
            source = names.get(int(edge["source_id"]), "未知卡片")
            target = names.get(int(edge["target_id"]), "未知卡片")
            relation = RELATION_LABELS.get(str(edge.get("relation")), "关联")
            document.add_paragraph(f"{source} → {target}（{relation}）", style="List Bullet")

    def _render_markdown(self, course: dict, notebook: dict, nodes: list[dict], edges: list[dict]) -> str:
        lines = [
            f"# {notebook['title']}",
            "",
            f"> 课程：{course['title']} · {len(nodes)} 张知识卡 · {len(edges)} 条关系",
            "",
        ]
        if notebook.get("description"):
            lines.extend([str(notebook["description"]), ""])
        lines.extend(["## 画布内容", ""])
        if not nodes:
            lines.extend(["这张画布暂时还没有知识卡片。", ""])
        for index, node in enumerate(nodes, start=1):
            lines.extend([
                f"### {index}. {node['title']}",
                "",
                f"`{KIND_LABELS.get(str(node.get('kind')), '知识卡')}` · {node.get('module') or '未分类'}",
                "",
            ])
            body = str(node.get("content") or node.get("description") or "").strip()
            if body:
                lines.extend([body, ""])
            image_path = self._image_path(node)
            if image_path:
                encoded = base64.b64encode(image_path.read_bytes()).decode("ascii")
                media_type = node.get("image_media_type") or "image/png"
                alt = str(node.get("image_alt") or node.get("title") or "图片")
                lines.extend([f"![{alt}](data:{media_type};base64,{encoded})", ""])
            if node.get("source_title"):
                lines.extend([f"**来源：** {node['source_title']}", ""])
            if node.get("source_quote"):
                quote = str(node["source_quote"]).replace("\n", "\n> ")
                lines.extend([f"> {quote}", ""])
        lines.extend(["## 关系清单", ""])
        names = {int(node["id"]): str(node["title"]) for node in nodes}
        if not edges:
            lines.extend(["尚未建立卡片关系。", ""])
        for edge in edges:
            source = names.get(int(edge["source_id"]), "未知卡片")
            target = names.get(int(edge["target_id"]), "未知卡片")
            relation = RELATION_LABELS.get(str(edge.get("relation")), "关联")
            lines.append(f"- {source} → {target}（{relation}）")
        lines.append("")
        return "\n".join(lines)

    def _image_path(self, node: dict) -> Path | None:
        storage_path = node.get("image_storage_path")
        if not storage_path:
            return None
        candidate = (self.media_root / str(storage_path)).resolve()
        if candidate.parent != self.media_root or not candidate.is_file():
            return None
        return candidate

    @staticmethod
    def _card_size(node: dict) -> tuple[int, int]:
        default_width = 280
        default_height = 310 if node.get("image_asset_id") else 190
        width = min(900, max(160, int(float(node.get("width") or default_width))))
        height = min(800, max(100, int(float(node.get("height") or default_height))))
        return (width, height)

    @staticmethod
    def _safe_filename(value: str) -> str:
        cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "-", str(value)).strip(" .")
        return (cleaned or "知识笔记")[:120]

    @staticmethod
    def _draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: tuple[int, int, int]) -> None:
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        length = 12
        spread = .52
        left = (end[0] - length * math.cos(angle - spread), end[1] - length * math.sin(angle - spread))
        right = (end[0] - length * math.cos(angle + spread), end[1] - length * math.sin(angle + spread))
        draw.polygon((end, left, right), fill=color)

    @staticmethod
    def _font(size: int, bold: bool = False):
        candidates = [
            Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
            Path("C:/Windows/Fonts/simhei.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        ]
        for candidate in candidates:
            if candidate.is_file():
                return ImageFont.truetype(str(candidate), size)
        return ImageFont.load_default()

    @staticmethod
    def _text_width(draw: ImageDraw.ImageDraw, value: str, font) -> int:
        return int(draw.textlength(value, font=font))

    @classmethod
    def _wrap(cls, draw: ImageDraw.ImageDraw, value: str, font, max_width: int) -> list[str]:
        lines: list[str] = []
        for paragraph in (value or "").splitlines() or [""]:
            current = ""
            for character in paragraph:
                candidate = current + character
                if current and cls._text_width(draw, candidate, font) > max_width:
                    lines.append(current)
                    current = character
                else:
                    current = candidate
            if current:
                lines.append(current)
        return lines
